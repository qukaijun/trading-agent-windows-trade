from __future__ import annotations

import re
from io import BytesIO
from typing import Any


def export_report_docx(markdown: str) -> bytes:
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("缺少 python-docx，无法导出 Word。") from exc

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    _set_east_asia_font(normal, "Microsoft YaHei", qn)

    for block in _parse_markdown(markdown):
        kind = block["kind"]
        text = block["text"]
        if kind == "title":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "Microsoft YaHei"
            _set_run_east_asia_font(run, "Microsoft YaHei", qn)
            paragraph.space_after = Pt(10)
        elif kind == "heading":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Microsoft YaHei"
            _set_run_east_asia_font(run, "Microsoft YaHei", qn)
            paragraph.space_before = Pt(8)
            paragraph.space_after = Pt(6)
        elif kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(text)
            run.font.name = "Microsoft YaHei"
            _set_run_east_asia_font(run, "Microsoft YaHei", qn)
        else:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = "Microsoft YaHei"
            _set_run_east_asia_font(run, "Microsoft YaHei", qn)
            paragraph.space_after = Pt(6)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_report_pdf(markdown: str) -> bytes:
    try:
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("缺少 reportlab，无法导出 PDF。") from exc

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "ChineseBody",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=16,
        alignment=TA_LEFT,
        spaceAfter=5,
    )
    title = ParagraphStyle(
        "ChineseTitle",
        parent=base,
        fontSize=18,
        leading=24,
        spaceAfter=10,
    )
    heading = ParagraphStyle(
        "ChineseHeading",
        parent=base,
        fontSize=14,
        leading=20,
        spaceBefore=8,
        spaceAfter=6,
    )

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story: list[Any] = []
    for block in _parse_markdown(markdown):
        kind = block["kind"]
        text = _escape_pdf_text(block["text"])
        if kind == "title":
            story.append(Paragraph(text, title))
        elif kind == "heading":
            story.append(Paragraph(text, heading))
        elif kind == "bullet":
            story.append(Paragraph(f"- {text}", base))
        else:
            story.append(Paragraph(text, base))
        story.append(Spacer(1, 2))

    doc.build(story)
    return buffer.getvalue()


def export_filename(report: dict[str, Any], suffix: str) -> str:
    symbol = report.get("symbol") or {}
    canonical = symbol.get("canonical") or symbol.get("raw") or "report"
    date = report.get("analysis_date") or "date"
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "-", f"{date}-{canonical}").strip("-")
    return f"{safe}.{suffix}"


def _parse_markdown(markdown: str) -> list[dict[str, str]]:
    blocks: list[dict[str, str]] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        text = _clean_inline(" ".join(line.strip() for line in paragraph_lines if line.strip()))
        if text:
            blocks.append({"kind": "paragraph", "text": text})
        paragraph_lines.clear()

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue
        if stripped.startswith("# "):
            flush_paragraph()
            blocks.append({"kind": "title", "text": _clean_inline(stripped[2:].strip())})
        elif stripped.startswith("## "):
            flush_paragraph()
            blocks.append({"kind": "heading", "text": _clean_inline(stripped[3:].strip())})
        elif stripped.startswith("- "):
            flush_paragraph()
            blocks.append({"kind": "bullet", "text": _clean_inline(stripped[2:].strip())})
        else:
            paragraph_lines.append(stripped)
    flush_paragraph()
    return blocks


def _clean_inline(text: str) -> str:
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    cleaned = re.sub(r"`([^`]*)`", r"\1", cleaned)
    return cleaned.strip()


def _escape_pdf_text(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def _set_east_asia_font(style: Any, font_name: str, qn: Any) -> None:
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_run_east_asia_font(run: Any, font_name: str, qn: Any) -> None:
    from docx.oxml import OxmlElement

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
